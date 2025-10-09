#!/usr/bin/env python3
"""
Import history questions from Word document with answer key
"""

import os
import json
import re
import psycopg2
from docx import Document

def extract_answer_key(doc):
    """Extract answer key from document"""
    all_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
    
    # Find where answers start
    answer_section = []
    in_answers = False
    
    for text in all_text:
        if 'Answers' in text or 'answers' in text:
            in_answers = True
            continue
        if in_answers:
            answer_section.append(text)
    
    # Parse answers like "1.d 2.a 3.a" etc.
    answers = {}
    
    for line in answer_section:
        # Find all patterns like "1.d" or "1.a" etc.
        matches = re.findall(r'(\d+)\.([a-dA-D])', line)
        for match in matches:
            q_num = int(match[0])
            answer_letter = match[1].lower()
            # Convert letter to index (a=0, b=1, c=2, d=3)
            answer_index = ord(answer_letter) - ord('a')
            answers[q_num] = answer_index
    
    return answers

def extract_questions_from_docx(docx_path):
    """Extract questions from Word document"""
    doc = Document(docx_path)
    
    # First extract the answer key
    print("📝 Extracting answer key...")
    answers = extract_answer_key(doc)
    print(f"✅ Found answers for {len(answers)} questions")
    
    # Extract questions
    all_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and not text.startswith('Answers') and '©' not in text and 'Reserved' not in text:
            # Stop before answer section
            if re.match(r'^\d+\.[a-d]\s+\d+\.[a-d]', text):
                break
            all_text.append(text)
    
    full_text = '\n'.join(all_text)
    
    questions = []
    current_question = None
    current_options = []
    question_num = 0
    
    for line in all_text:
        line = line.strip()
        if not line:
            continue
        
        # Check if this is a question start (number followed by dot or parenthesis)
        q_match = re.match(r'^(\d+)[\.\)]\s*(.+)', line)
        if q_match:
            # Save previous question if exists
            if current_question and len(current_options) >= 4:
                q_num = question_num
                if q_num in answers:
                    questions.append({
                        'question': current_question,
                        'options': current_options[:4],
                        'correct_answer': answers[q_num]
                    })
            
            # Start new question
            question_num = int(q_match.group(1))
            current_question = q_match.group(2).strip()
            current_question = re.sub(r'…+', '', current_question).strip()
            current_options = []
            continue
        
        # Check if this is an option line (starts with a), b), c), d))
        opt_match = re.match(r'^\s*([a-d])\)\s*(.+)', line, re.IGNORECASE)
        if opt_match and current_question:
            option_text = opt_match.group(2).strip()
            # Remove any trailing answer indicators
            option_text = re.sub(r'\s+[a-d]\).*$', '', option_text, flags=re.IGNORECASE)
            current_options.append(option_text)
            continue
        
        # Check if line contains multiple options on same line
        if current_question and re.search(r'[a-d]\)', line, re.IGNORECASE):
            # Split by option markers
            parts = re.split(r'\s+([a-d])\)\s*', line, flags=re.IGNORECASE)
            for i in range(1, len(parts), 2):
                if i+1 < len(parts):
                    option_text = parts[i+1].strip()
                    # Clean up option text
                    option_text = re.sub(r'\s+[a-d]\).*$', '', option_text, flags=re.IGNORECASE)
                    option_text = option_text.split('\t')[0].strip()
                    if option_text:
                        current_options.append(option_text)
            continue
        
        # Otherwise, append to current question text
        if current_question and not current_options:
            current_question += ' ' + line.strip()
            current_question = re.sub(r'…+', '', current_question).strip()
    
    # Save last question
    if current_question and len(current_options) >= 4 and question_num in answers:
        questions.append({
            'question': current_question,
            'options': current_options[:4],
            'correct_answer': answers[question_num]
        })
    
    return questions

def import_to_database(questions):
    """Import questions to PostgreSQL database"""
    database_url = os.environ.get('DATABASE_URL')
    if not database_url:
        print("❌ DATABASE_URL environment variable not set")
        return False
    
    try:
        conn = psycopg2.connect(database_url)
        cursor = conn.cursor()
        
        # Get current question count
        cursor.execute("SELECT COUNT(*) FROM questions")
        current_count = cursor.fetchone()[0]
        print(f"\n📊 Current questions in database: {current_count}")
        
        imported = 0
        skipped = 0
        
        for q in questions:
            # Skip if less than 4 options
            if len(q['options']) < 4:
                skipped += 1
                continue
            
            # Clean question text
            question_text = q['question'].strip()
            if not question_text or len(question_text) < 10:
                skipped += 1
                continue
            
            # Convert options to proper JSON
            options_json = json.dumps(q['options'])
            
            try:
                cursor.execute(
                    "INSERT INTO questions (question, options, correct_answer) VALUES (%s, %s, %s)",
                    (question_text, options_json, q['correct_answer'])
                )
                imported += 1
                
                if imported % 50 == 0:
                    print(f"  ⏳ Imported {imported} questions...")
                    
            except psycopg2.IntegrityError:
                # Duplicate question, skip
                skipped += 1
                conn.rollback()
                continue
        
        conn.commit()
        
        # Get new total
        cursor.execute("SELECT COUNT(*) FROM questions")
        new_count = cursor.fetchone()[0]
        
        cursor.close()
        conn.close()
        
        print(f"\n✅ Successfully imported {imported} new questions!")
        print(f"⚠️  Skipped {skipped} questions (duplicates or invalid format)")
        print(f"📊 Total questions in database: {new_count}")
        
        return True
        
    except Exception as e:
        print(f"❌ Database error: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("=" * 70)
    print("📚 Importing History Questions from Word Document")
    print("=" * 70)
    print()
    
    docx_path = 'attached_assets/5_6284844636582190663_1760023077830.docx'
    
    print("📖 Reading Word document...")
    questions = extract_questions_from_docx(docx_path)
    
    print(f"\n✅ Extracted {len(questions)} questions from document")
    
    # Show first 5 questions as preview
    print("\n" + "=" * 70)
    print("📋 Preview of first 5 questions:")
    print("=" * 70)
    for i, q in enumerate(questions[:5], 1):
        print(f"\n{i}. {q['question']}")
        for j, opt in enumerate(q['options']):
            marker = "✓" if j == q['correct_answer'] else " "
            print(f"   {chr(65+j)}) {opt} {marker}")
    
    print("\n" + "=" * 70)
    print("🗄️  Starting import to database...")
    print("=" * 70)
    import_to_database(questions)
